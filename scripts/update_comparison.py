"""
Rebuild comparison section - software-only, subscription model justified by
bundled forecasting + Solcast integration that competitors charge separately.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from _docx_utils import (
    collect_siblings_until,
    find_first_paragraph,
    find_next_numbered_heading,
    is_paragraph,
    paragraph_text,
    remove_elements,
    resolve_input_output,
)

in_path, out_path = resolve_input_output("Rebuild Industry Comparison section.")
doc = Document(str(in_path))
body = doc.element.body

def set_cell(cell, text, bold=False, size=9, color=None):
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, hex_color):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(s)

def make_paragraph(text, size=10, bold=False, color=None):
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), str(120))
    pPr.append(spacing)
    new_p.append(pPr)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(size * 2))
    rPr.append(szCs)
    font_el = OxmlElement("w:rFonts")
    font_el.set(qn("w:ascii"), "Calibri")
    font_el.set(qn("w:hAnsi"), "Calibri")
    rPr.append(font_el)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "{:02x}{:02x}{:02x}".format(*color))
        rPr.append(c)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    new_p.append(run)
    return new_p

# Locate the comparison heading by text (03/04 compatible)
comparison_heading = find_first_paragraph(
    doc, lambda t: "Industry Comparison" in (t or "")
)
if comparison_heading is None:
    raise RuntimeError("Comparison heading not found (expected text containing 'Industry Comparison').")

# End boundary: "Prepared by"/"ENGR." paragraph if present, else next numbered heading.
stop_el = None
cur = comparison_heading.getnext()
while cur is not None:
    if is_paragraph(cur):
        txt = paragraph_text(cur).strip()
        txt_lower = txt.lower()
        if "prepared by" in txt_lower or "engr." in txt_lower:
            stop_el = cur
            break
    cur = cur.getnext()

if stop_el is None:
    stop_el = find_next_numbered_heading(comparison_heading)

to_remove = collect_siblings_until(comparison_heading, stop_el)
if to_remove:
    remove_elements(body, to_remove)

# SOFTWARE-ONLY comparison
# Forecasting row now split into two clear rows to emphasize bundled value
# Cost row reframed: ADSI subscription includes forecasting, others do not
features = [
    ["Feature", "ADSI Dashboard", "Survalent\nStationCentral\n(NTA Quote)", "Ignition SCADA", "GreenPower-\nMonitor (DNV)", "Huawei\nFusionSolar", "Custom SCADA"],
    ["Real-Time Monitoring", "Yes - live polling,\nWebSocket updates", "Yes - DNP3 /\nModbus polling", "Yes - OPC-UA\nbased", "Yes -\ncloud-based", "Yes -\nHuawei only", "Varies"],
    ["Inverter Start/Stop\nControl", "Yes - per-node,\nper-inverter, bulk", "Yes - via\ncontrol commands", "Yes (driver dev\nrequired)", "Limited", "Huawei inverters\nonly", "Varies"],
    ["Plant Output\nCapping", "Yes - built-in\nauto controller", "No", "Custom scripting\nneeded", "No", "No", "Custom dev\nneeded"],
    ["ML Day-Ahead\nForecasting", "Included\nin subscription", "Not available\n(separate purchase\nrequired)", "Not available\n(3rd party needed)", "Optional\nadd-on cost", "Not available", "Custom dev\nneeded"],
    ["Solcast Solar\nForecasting", "Included\nin subscription", "Not available\n(separate Solcast\nsubscription needed)", "Not available\n(separate Solcast\nsubscription needed)", "Separate\ncontract", "Not available", "Separate Solcast\nsubscription needed"],
    ["Intraday Adjusted\nForecast", "Included\nin subscription", "Not available", "Not available", "Optional", "Not available", "Custom dev\nneeded"],
    ["Alarm Management", "Yes - episodes,\naudio, audit trail", "Yes - alarm\narchive retentive", "Yes - advanced", "Yes", "Basic", "Varies"],
    ["Daily Reports &\nAnalytics", "Yes - automated,\nexportable", "Limited -\nHMI display only", "Yes\n(custom dev)", "Yes", "Basic", "Custom dev\nneeded"],
    ["Data Export\n(CSV/XLSX)", "Yes - all\ndata types", "Limited", "Yes", "Yes", "Limited", "Varies"],
    ["Dynamic Availability\nCalculation", "Yes - per-inverter,\nauto window", "No", "Custom\nscripting", "Yes", "Basic", "Custom dev\nneeded"],
    ["Cloud Backup", "Yes - OneDrive,\nGoogle Drive", "No", "No\n(separate infra)", "Cloud-native", "Huawei Cloud\nonly", "Separate\ninfra"],
    ["Gateway-Remote\nArchitecture", "Yes - built-in,\nno extra license", "No - single\nserver license", "Extra licenses\nneeded", "Cloud\nmulti-user", "Cloud portal", "Custom dev\nneeded"],
    ["Data Replication\n& Sync", "Yes - incremental,\narchive-aware", "No", "Module\nrequired", "Cloud-native", "Cloud-only", "Custom dev\nneeded"],
    ["Operator\nMessaging", "Yes - built-in\nchat channel", "No", "No", "No", "No", "Custom dev\nneeded"],
    ["Multi-Brand\nInverter Support", "Modbus TCP\n(adaptable)", "Modbus + DNP3\n(multi-protocol)", "OPC-UA /\ndrivers", "Multi-vendor", "Huawei only", "Varies"],
    ["Auto-Update\nSystem", "Yes - built-in\nupdater", "No", "Manual", "Cloud-managed", "Cloud-managed", "Manual"],
    ["Vendor Lock-In", "None", "Medium\n(NTA/Survalent)", "Low", "Medium", "High\n(Huawei only)", "None"],
    ["Local Support\n(Philippines)", "Yes - direct\ndeveloper access", "Yes - NTA\nlocal office", "Foreign vendor", "Foreign vendor", "Foreign vendor", "Varies"],
    ["Software License\nCost", "From Php 55,000/mo\n(all-inclusive)", "Php 850,000\nperpetual\n(Ref: 2023-J-081)", "USD 15,000 -\n35,000+\nperpetual", "USD 30,000 -\n80,000+\nannual", "Free\n(Huawei only)", "USD 50,000 -\n200,000+"],
    ["Forecasting\nIncluded in License", "Yes - ML engine\n+ Solcast bundled", "No - additional\ncost required", "No - additional\ncost required", "No - separate\ncontract", "No", "No - additional\ncost required"],
]

elements_to_insert = []

intro_p = make_paragraph(
    "The following comparison evaluates the ADSI Inverter Dashboard software against "
    "industry-standard SCADA and solar monitoring platforms across key software features, "
    "integration capabilities, and licensing cost. The Survalent StationCentral pricing is "
    "based on an actual quotation (Ref: 2023-J-081 R1) received for the same plant site.",
    size=10,
)
elements_to_insert.append(intro_p)
elements_to_insert.append(make_paragraph("", size=6))

# Feature Table - 7 columns
NUM_COLS = 7
table = doc.add_table(rows=len(features), cols=NUM_COLS)
table.autofit = True
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
borders = OxmlElement("w:tblBorders")
for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "4")
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "999999")
    borders.append(el)
tblPr.append(borders)

header_bg = "1a1a2e"
for j, h in enumerate(features[0]):
    set_cell(table.rows[0].cells[j], h, bold=True, size=7.5, color=(255, 255, 255))
    shade_cell(table.rows[0].cells[j], header_bg)

for i in range(1, len(features)):
    for j in range(NUM_COLS):
        txt = features[i][j]
        is_feature = j == 0
        c = None
        if j == 1 and (txt.startswith("Yes") or txt.startswith("Included")):
            c = (0x0A, 0x7E, 0x2A)
        elif j > 1 and (txt.startswith("No") or txt.startswith("Not available")):
            c = (0xAA, 0x33, 0x33)
        set_cell(table.rows[i].cells[j], txt, bold=is_feature, size=7, color=c)

    if i % 2 == 0:
        for j in range(NUM_COLS):
            shade_cell(table.rows[i].cells[j], "f0f0f8")

elements_to_insert.append(table._element)
elements_to_insert.append(make_paragraph("", size=6))

# Subscription value note
elements_to_insert.append(make_paragraph(
    "Why Subscription?",
    size=11, bold=True, color=(0x1A, 0x1A, 0x2E),
))
elements_to_insert.append(make_paragraph("", size=2))

sub_notes = [
    "The ADSI Dashboard subscription model bundles ML-based day-ahead forecasting, "
    "Solcast solar forecasting integration, intraday forecast adjustment, continuous "
    "software updates, and full technical support into a single monthly fee.",
    "With other platforms, forecasting requires a separate Solcast API subscription "
    "or third-party service contract on top of the SCADA license. "
    "The ADSI subscription eliminates that additional cost and consolidates "
    "monitoring, control, and forecasting into one provider.",
]
for note in sub_notes:
    elements_to_insert.append(make_paragraph(note, size=10))

elements_to_insert.append(make_paragraph("", size=6))

# Key Advantages heading
elements_to_insert.append(make_paragraph(
    "Key Advantages of Choosing ADSI Inverter Dashboard:",
    size=11, bold=True, color=(0x1A, 0x1A, 0x2E),
))
elements_to_insert.append(make_paragraph("", size=4))

advantages = [
    "All-inclusive subscription - monitoring, control, forecasting, and Solcast integration bundled in one fee. No separate forecast service contracts.",
    "No large upfront capital expenditure. Survalent requires Php 850,000 perpetual license before any forecasting capability is even considered.",
    "Built-in ML forecasting engine and Solcast integration that competitors do not offer or charge separately for.",
    "Complete software solution: real-time monitoring, fast inverter control, plant output capping, alarms, audit trail, automated reports, data export, cloud backup, and remote operations.",
    "Gateway-Remote architecture supports multi-workstation deployment without additional licensing fees.",
    "No vendor lock-in - works with Ingeteam and adaptable to other inverter brands via Modbus TCP.",
    "Locally developed and supported by an in-house developer - faster response, zero foreign vendor escalation, aligned with Philippine utility operations.",
    "Continuous updates and new features delivered automatically through built-in updater at no additional cost.",
]

for adv in advantages:
    p = make_paragraph("- " + adv, size=10)
    pPr = p.find(qn("w:pPr"))
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    pPr.append(ind)
    elements_to_insert.append(p)

elements_to_insert.append(make_paragraph("", size=10))

# Insert all rebuilt elements after the comparison heading
ref = comparison_heading
for elem in elements_to_insert:
    ref.addnext(elem)
    ref = elem

doc.save(str(out_path))
print(f"DONE - saved successfully: {out_path}")
