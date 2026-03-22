"""
Expand Section 02 After-Market Rates with more detail,
and add a Perpetual License option between Section 01 and 02.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from _docx_utils import (
    collect_from_to,
    collect_siblings_until,
    find_first_paragraph,
    find_next_numbered_heading,
    remove_elements,
    resolve_input_output,
)

in_path, out_path = resolve_input_output("Update After-Market and Perpetual sections.")
doc = Document(str(in_path))
body = doc.element.body

def set_cell(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        aligns = {"center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
        p.alignment = aligns.get(align)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, hex_color):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(s)

def add_borders(table, color="999999"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)

def make_paragraph(text, size=10, bold=False, color=None):
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), str(120))
    pPr.append(spacing)
    new_p.append(pPr)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(size * 2))
    rPr.append(szCs)
    font_el = OxmlElement("w:rFonts")
    font_el.set(qn("w:ascii"), "Calibri")
    font_el.set(qn("w:hAnsi"), "Calibri")
    rPr.append(font_el)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "{:02x}{:02x}{:02x}".format(*color))
        rPr.append(c)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    new_p.append(run)
    return new_p

HEADER_BG = "1a1a2e"
DARK = (0x1A, 0x1A, 0x2E)
GREEN = (0x0A, 0x7E, 0x2A)
WHITE = (255, 255, 255)
ALT_ROW = "f0f0f8"

# ============================================================
# STEP 1: Find and replace Section 02 heading and its table
# ============================================================

sec02_elem = find_first_paragraph(
    doc,
    lambda t: (t or "").strip().startswith("02.") and "After-Market" in (t or ""),
) or find_first_paragraph(
    doc,
    lambda t: (t or "").strip().startswith("03.") and "After-Market" in (t or ""),
)
if sec02_elem is None:
    raise RuntimeError("After-Market section heading not found.")

comparison_elem = find_first_paragraph(
    doc, lambda t: "Industry Comparison" in (t or "")
)
if comparison_elem is None:
    comparison_elem = find_next_numbered_heading(sec02_elem)

# Remove existing Section 02 body content before rebuilding.
existing_sec02_body = collect_siblings_until(sec02_elem, comparison_elem)
if existing_sec02_body:
    remove_elements(body, existing_sec02_body)

elements_02 = []

# Intro
elements_02.append(make_paragraph(
    "The following services are available for work beyond the standard subscription scope. "
    "All after-market work is scoped and quoted before engagement.",
    size=10,
))
elements_02.append(make_paragraph("", size=4))

# Expanded After-Market table
aftermarket = [
    ["Service", "Scope", "Estimated Rate"],
    ["UI Customization",
     "Layout adjustments, branding,\ntheme modifications, custom\ndashboard views",
     "Php 25,000 - 40,000"],
    ["Additional Features\nor Functions",
     "New software modules, custom\nreports, additional export\nformats, workflow automation",
     "Php 30,000 - 60,000"],
    ["New Inverter Brand /\nProtocol Integration",
     "Modbus register mapping,\ndriver development, testing,\nand commissioning for a\nnew inverter manufacturer",
     "From Php 150,000"],
    ["Additional Site\nDeployment",
     "Installation, configuration,\nand commissioning of the\nplatform on a new plant site",
     "Quoted per site"],
    ["On-Site Technical\nSupport",
     "Physical site visit for\ntroubleshooting, commissioning,\nor training",
     "Quoted per visit\n(travel + daily rate)"],
    ["Custom Data\nIntegration",
     "Integration with third-party\nsystems (e.g., weather APIs,\ngrid dispatch, EMS/DCS)",
     "Php 40,000 - 100,000"],
]

tbl_am = doc.add_table(rows=len(aftermarket), cols=3)
tbl_am.autofit = True
add_borders(tbl_am)

for j in range(3):
    set_cell(tbl_am.rows[0].cells[j], aftermarket[0][j], bold=True, size=9, color=WHITE)
    shade_cell(tbl_am.rows[0].cells[j], HEADER_BG)

for i in range(1, len(aftermarket)):
    set_cell(tbl_am.rows[i].cells[0], aftermarket[i][0], bold=True, size=9, color=DARK)
    set_cell(tbl_am.rows[i].cells[1], aftermarket[i][1], size=8.5)
    set_cell(tbl_am.rows[i].cells[2], aftermarket[i][2], bold=True, size=9)
    if i % 2 == 0:
        for j in range(3):
            shade_cell(tbl_am.rows[i].cells[j], ALT_ROW)

elements_02.append(tbl_am._element)
elements_02.append(make_paragraph("", size=6))

# Insert after sec02 heading
ref = sec02_elem
for elem in elements_02:
    ref.addnext(elem)
    ref = elem

# ============================================================
# STEP 2: Insert Perpetual License section between 01 and 02
# ============================================================

# Remove existing 01-B block if present (for idempotent reruns).
existing_01b = find_first_paragraph(
    doc, lambda t: "01-B. Perpetual License Option" in (t or "")
)
if existing_01b is not None:
    # Only remove if the existing 01-B block is before the current After-Market heading.
    probe = existing_01b
    seen_sec02 = False
    while probe is not None:
        if probe is sec02_elem:
            seen_sec02 = True
            break
        probe = probe.getnext()
    if seen_sec02:
        old_perpetual_block = collect_from_to(existing_01b, sec02_elem)
        if old_perpetual_block:
            remove_elements(body, old_perpetual_block)

# Build perpetual section elements - insert BEFORE sec02
perp_elements = []

# Section heading
perp_elements.append(make_paragraph(
    "01-B. Perpetual License Option",
    size=14, bold=True, color=DARK,
))
perp_elements.append(make_paragraph("", size=4))
perp_elements.append(make_paragraph(
    "For organizations that prefer a one-time capital expenditure over recurring subscription fees, "
    "a perpetual license is available. The perpetual license includes the same complete software "
    "platform with all features. Annual support renewal is optional after the first year.",
    size=10,
))
perp_elements.append(make_paragraph("", size=4))

# Perpetual pricing table
perp_data = [
    ["Item", "Price", "Details"],
    ["Perpetual Software\nLicense",
     "Php 1,500,000\none-time",
     "Full ownership of the current software version.\n"
     "Includes all features: monitoring, control,\n"
     "forecasting (ML + Solcast), reports, exports,\n"
     "cloud backup, and remote operations.\n"
     "First year of support and updates included."],
    ["Annual Support\nRenewal (Optional)",
     "Php 180,000 / yr.",
     "Continued software updates, new features,\n"
     "remote technical support, and Solcast\n"
     "integration maintenance.\n"
     "Without renewal, software continues to work\n"
     "but will not receive updates or support."],
]

tbl_perp = doc.add_table(rows=len(perp_data), cols=3)
tbl_perp.autofit = True
add_borders(tbl_perp)

for j in range(3):
    set_cell(tbl_perp.rows[0].cells[j], perp_data[0][j], bold=True, size=9, color=WHITE)
    shade_cell(tbl_perp.rows[0].cells[j], HEADER_BG)

for i in range(1, len(perp_data)):
    set_cell(tbl_perp.rows[i].cells[0], perp_data[i][0], bold=True, size=9, color=DARK)
    set_cell(tbl_perp.rows[i].cells[1], perp_data[i][1], bold=True, size=9)
    set_cell(tbl_perp.rows[i].cells[2], perp_data[i][2], size=8.5)
    if i % 2 == 0:
        for j in range(3):
            shade_cell(tbl_perp.rows[i].cells[j], ALT_ROW)

perp_elements.append(tbl_perp._element)
perp_elements.append(make_paragraph("", size=4))

# Comparison note
perp_elements.append(make_paragraph(
    "Perpetual vs Subscription at a Glance",
    size=10, bold=True, color=DARK,
))

compare = [
    ["", "Perpetual License", "Subscription (5-Year)"],
    ["Upfront Cost", "Php 1,500,000", "None"],
    ["Monthly Cost", "None\n(after purchase)", "Php 55,000 / mo."],
    ["5-Year Total Cost\n(with annual support)", "Php 2,220,000\n(license + 4 yrs support)", "Php 3,100,000\n(upfront) or\nPhp 3,300,000\n(monthly)"],
    ["Software Ownership", "Permanent", "Active subscription\nrequired"],
    ["Updates & Support\nAfter Year 1", "Optional renewal\n(Php 180,000/yr)", "Included in\nsubscription"],
    ["Forecasting\n(ML + Solcast)", "Included", "Included"],
]

tbl_cmp = doc.add_table(rows=len(compare), cols=3)
tbl_cmp.autofit = True
add_borders(tbl_cmp)

for j in range(3):
    set_cell(tbl_cmp.rows[0].cells[j], compare[0][j], bold=True, size=9, color=WHITE)
    shade_cell(tbl_cmp.rows[0].cells[j], HEADER_BG)

for i in range(1, len(compare)):
    set_cell(tbl_cmp.rows[i].cells[0], compare[i][0], bold=True, size=9, color=DARK)
    for j in range(1, 3):
        set_cell(tbl_cmp.rows[i].cells[j], compare[i][j], size=9, align="center")
    if i % 2 == 0:
        for j in range(3):
            shade_cell(tbl_cmp.rows[i].cells[j], ALT_ROW)

perp_elements.append(tbl_cmp._element)
perp_elements.append(make_paragraph("", size=8))

# Insert perpetual section BEFORE sec02 heading (in reverse so order is correct)
for elem in reversed(perp_elements):
    sec02_elem.addprevious(elem)

# ============================================================
# STEP 3: Renumber Section 02 -> 03, Section 03 -> 04
# ============================================================
for p in doc.paragraphs:
    txt = (p.text or "").strip()
    if txt.startswith(("02.", "03.", "04.")) and "After-Market" in txt:
        p.text = ""
        run = p.add_run("03. After-Market Rates")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*DARK)
        break

for p in doc.paragraphs:
    txt = (p.text or "").strip()
    if txt.startswith(("03.", "04.", "05.")) and "Industry Comparison" in txt:
        p.text = ""
        run = p.add_run("04. Industry Comparison: Features, Integrations, and Pricing")
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*DARK)
        break

doc.save(str(out_path))
print(f"DONE - saved successfully: {out_path}")
