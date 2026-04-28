"""
Break down Section 01 pricing table into a more informative layout:
- Table A: What's Included (feature breakdown)
- Table B: Subscription Plans with clear tier comparison
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from _docx_utils import (
    collect_siblings_until,
    find_first_paragraph,
    find_next_numbered_heading,
    remove_elements,
    resolve_input_output,
)

in_path, out_path = resolve_input_output("Update Section 01 pricing tables.")
doc = Document(str(in_path))
body = doc.element.body

def set_cell(cell, text, bold=False, size=9, color=None, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        aligns = {"center": WD_ALIGN_PARAGRAPH.CENTER, "right": WD_ALIGN_PARAGRAPH.RIGHT}
        p.alignment = aligns.get(align)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = RGBColor(*color)

def shade_cell(cell, hex_color):
    shading = cell._element.get_or_add_tcPr()
    s = shading.makeelement(qn("w:shd"), {
        qn("w:val"): "clear",
        qn("w:color"): "auto",
        qn("w:fill"): hex_color,
    })
    shading.append(s)

def add_borders(table, color="999999"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)

def make_paragraph(text, size=10, bold=False, color=None):
    new_p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), str(120))
    pPr.append(spacing)
    new_p.append(pPr)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), str(size * 2))
    rPr.append(szCs)
    font_el = OxmlElement("w:rFonts")
    font_el.set(qn("w:ascii"), "Calibri")
    font_el.set(qn("w:hAnsi"), "Calibri")
    rPr.append(font_el)
    if bold:
        b = OxmlElement("w:b")
        rPr.append(b)
    if color:
        c = OxmlElement("w:color")
        c.set(qn("w:val"), "{:02x}{:02x}{:02x}".format(*color))
        rPr.append(c)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run.append(t)
    new_p.append(run)
    return new_p

# --- Remove existing Section 01 body content by heading anchors ---
sec01_elem = find_first_paragraph(
    doc, lambda t: (t or "").strip().startswith("01. Proposed Pricing Summary")
)
if sec01_elem is None:
    raise RuntimeError("Section heading '01. Proposed Pricing Summary' not found.")

next_heading = find_next_numbered_heading(sec01_elem)
if next_heading is None:
    raise RuntimeError("Could not find next numbered heading after section 01.")

to_remove = collect_siblings_until(sec01_elem, next_heading)
if to_remove:
    remove_elements(body, to_remove)

# --- Build new content ---
elements = []

# Intro note
elements.append(make_paragraph(
    "All subscription plans include the complete ADSI Inverter Dashboard software platform "
    "with full access to every feature listed below. Plans differ only in commitment length and pricing.",
    size=10,
))
elements.append(make_paragraph("", size=4))

# --- Table A: What's Included in Every Plan ---
elements.append(make_paragraph(
    "What's Included",
    size=11, bold=True, color=(0x1A, 0x1A, 0x2E),
))

inclusions = [
    ["Category", "Features Included"],
    ["Monitoring & Control",
     "Real-time inverter monitoring (live polling + WebSocket)\n"
     "Per-node, per-inverter, and bulk Start/Stop control\n"
     "Plant output capping with automatic controller"],
    ["Forecasting & Analytics",
     "ML-based day-ahead power forecasting (built-in engine)\n"
     "Solcast solar forecasting integration (no separate subscription)\n"
     "Intraday adjusted forecast with auto-correction\n"
     "Dynamic per-inverter availability calculation"],
    ["Reporting & Export",
     "Automated daily performance reports\n"
     "CSV and XLSX data export for all data types\n"
     "Forecast vs actual comparison reports"],
    ["Alarms & Audit",
     "Alarm episode management with audio alerts\n"
     "Full audit trail logging\n"
     "Alarm acknowledgment and history"],
    ["Data Management",
     "Cloud backup (OneDrive / Google Drive)\n"
     "Incremental data replication and archive-aware sync\n"
     "Hot/cold data architecture with automatic archival"],
    ["Remote Operations",
     "Gateway-Remote architecture (no extra license)\n"
     "Multi-workstation access over LAN or Tailscale VPN\n"
     "Operator messaging (gateway-to-remote chat)"],
    ["Software Maintenance",
     "Automatic software updates via built-in updater\n"
     "Bug fixes, security patches, and feature enhancements\n"
     "Remote technical support from the developer"],
]

tbl_a = doc.add_table(rows=len(inclusions), cols=2)
tbl_a.autofit = True
add_borders(tbl_a)

header_bg = "1a1a2e"
for j in range(2):
    set_cell(tbl_a.rows[0].cells[j], inclusions[0][j], bold=True, size=9, color=(255, 255, 255))
    shade_cell(tbl_a.rows[0].cells[j], header_bg)

for i in range(1, len(inclusions)):
    set_cell(tbl_a.rows[i].cells[0], inclusions[i][0], bold=True, size=9, color=(0x1A, 0x1A, 0x2E))
    set_cell(tbl_a.rows[i].cells[1], inclusions[i][1], size=8.5)
    if i % 2 == 0:
        shade_cell(tbl_a.rows[i].cells[0], "f0f0f8")
        shade_cell(tbl_a.rows[i].cells[1], "f0f0f8")

elements.append(tbl_a._element)
elements.append(make_paragraph("", size=6))

# --- Table B: Subscription Plans ---
elements.append(make_paragraph(
    "Subscription Plans",
    size=11, bold=True, color=(0x1A, 0x1A, 0x2E),
))

plans = [
    ["", "1-Year Plan", "3-Year Plan", "5-Year Plan"],
    ["Monthly Rate", "Php 70,000 / mo.", "Php 60,000 / mo.", "Php 55,000 / mo."],
    ["Annual / Term Cost", "Php 760,000 / yr.\n(save Php 80,000\nvs monthly)", "Php 2,000,000 / 3 yrs.\n(save Php 160,000\nvs monthly)", "Php 3,100,000 / 5 yrs.\n(save Php 200,000\nvs monthly)"],
    ["Billing Options", "Monthly or\nAnnual", "Monthly or\nUpfront", "Monthly or\nUpfront"],
    ["Software Access", "Full access to\nall features", "Full access to\nall features", "Full access to\nall features"],
    ["ML Forecasting\n& Solcast", "Included", "Included", "Included"],
    ["Software Updates\n& Enhancements", "Included", "Included", "Included"],
    ["Remote Support", "Included", "Included", "Priority support"],
    ["Dedicated Account\nHandling", "-", "-", "Included"],
    ["Rate Lock", "Current term only", "Locked for 3 years", "Locked for 5 years"],
]

tbl_b = doc.add_table(rows=len(plans), cols=4)
tbl_b.autofit = True
add_borders(tbl_b)

for j in range(4):
    set_cell(tbl_b.rows[0].cells[j], plans[0][j], bold=True, size=9, color=(255, 255, 255))
    shade_cell(tbl_b.rows[0].cells[j], header_bg)

for i in range(1, len(plans)):
    set_cell(tbl_b.rows[i].cells[0], plans[i][0], bold=True, size=9, color=(0x1A, 0x1A, 0x2E))
    for j in range(1, 4):
        is_included = plans[i][j] in ("Included", "Priority support")
        c = (0x0A, 0x7E, 0x2A) if is_included else None
        set_cell(tbl_b.rows[i].cells[j], plans[i][j], size=9, color=c, align="center")
    if i % 2 == 0:
        for j in range(4):
            shade_cell(tbl_b.rows[i].cells[j], "f0f0f8")

elements.append(tbl_b._element)
elements.append(make_paragraph("", size=4))

# Note about after-market
elements.append(make_paragraph(
    "Customization, additional features, or new inverter/protocol integrations beyond "
    "the subscription scope are quoted separately under After-Market Services (see Section 02).",
    size=9,
))
elements.append(make_paragraph("", size=6))

# Insert all elements after Section 01 heading
ref = sec01_elem
for elem in elements:
    ref.addnext(elem)
    ref = elem

doc.save(str(out_path))
print(f"DONE - saved successfully: {out_path}")
